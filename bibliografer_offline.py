from html import unescape
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm

class Exceptions:

	def signal_exception(reason):
		
		print(reason)
		quit()

class Philpapers:

	def check_urls(raw_urls):
	
		urls_count = raw_urls.count('https://philpapers.org/rec/')
		urls_processing = raw_urls
		urls_list = []
		
		if urls_count == 0:
			exception_reason = "Brak poprawnych adresów URL odnoszących do Philpapers"
			Exceptions.signal_exception(exception_reason)
		
		for i in range(urls_count):
			nth_url = urls_processing[urls_processing.rfind('https://philpapers.org/rec/'):]
			urls_processing = urls_processing[0:len(urls_processing)-len(nth_url)]
			nth_url = nth_url.strip(', ')
			urls_list.append(nth_url)
			
		urls_list[:] = list(set(urls_list))

		return urls_list
	
	def get_data(philpapers_url):
	
		r_lines_dict = []
		
		print("ppapers open")
		r = requests.get(philpapers_url, stream=True)

		_delimiter = 'Options'

		for line in r.iter_lines():
			_line = line.decode('utf-8','ignore')
			if _line != _delimiter:
				r_lines_dict.append(_line)
			else:
				break
				
		r.close()
		print("closed")
		
		r_lines_str = ''.join(r_lines_dict)
		
		return r_lines_str
	
	def check_integrity(raw_data, raw_data_url):
	
		# if raw_data.find('schema.org/Article') > 0:
			# Philpapers.paper_is_article = True
		# elif raw_data.find('schema.org/Book') > 0:
			# Philpapers.paper_is_book = True
		if raw_data.count('schema.org/Book') == 0 and raw_data.count('schema.org/Article') == 0:
			integrity = False
		
		return integrity
	
	def parse_data(raw_data):
		
		paper_is_book = False
		paper_is_article = False
		article_in_book = False
		article_in_journal = False
		paper_type = None
		paper_data = None
		
		# ten czek łatwiej zrobić jak jeszcze wszystko jest stringiem
		
		if raw_data.find('schema.org/Article') > 0:
			paper_is_article = True
		elif raw_data.find('schema.org/Book') > 0:
			paper_is_book = True
		
		r_bs = BeautifulSoup(raw_data, 'html.parser')
		
		# to się dzieje w meta:
		
		bs_authors_list = r_bs.find_all(attrs={'name':'citation_author'})
		authors_number = len(bs_authors_list)
		
		bs_title = str(r_bs.find_all(attrs={'name':'citation_title'}))
		bs_date = str(r_bs.find_all(attrs={'name':'citation_publication_date'}))
		
		paper_title = bs_title[bs_title.find('content=')+9:bs_title.rfind('name=')-2]
		paper_date = bs_date[bs_date.find('content=')+9:bs_date.rfind('name=')-2]
		
		if authors_number == 1:
			temp_author = str(bs_authors_list[0])
			temp_author = temp_author[temp_author.find('content=')+9:temp_author.rfind('name=')-2]
			temp_author_lastname = temp_author[temp_author.rfind(' '):].strip()
			temp_author_firstname = temp_author[:temp_author.rfind(' ')].strip()
			paper_authors = temp_author_lastname + ' ' + temp_author_firstname
		else:
			temp_authors_list = []
			for item in bs_authors_list:
				temp_author = str(item)
				temp_author = temp_author[temp_author.find('content=')+9:temp_author.rfind('name=')-2]
				temp_author_lastname = temp_author[temp_author.rfind(' '):].strip()
				temp_author_firstname = temp_author[:temp_author.rfind(' ')].strip()
				temp_author = temp_author_lastname + ' ' + temp_author_firstname
				temp_authors_list.append(temp_author)
			paper_authors = ', '.join(temp_authors_list)
		
		# to się dzieje w recAuthors:
		
		bs_recauthors = str(r_bs.find_all(class_='recAuthors'))
		
		if bs_recauthors.count('(ed.)') == 1:
			paper_authors = paper_authors + ' (ed.)'
		elif bs_recauthors.count('(eds.)') == 1:
			paper_authors = paper_authors + ' (eds.)'
			
		# to się dzieje w pubinfo:
				
		bs_pubinfo = str(r_bs.find_all(class_='recPubInfo'))
		
		if bs_pubinfo.find('asearch.pl') > 0:
			article_in_journal = True
		else:
			article_in_book = True
			
		if paper_is_book:
			paper_publisher = bs_pubinfo[bs_pubinfo.find('recPubInfo')+12:bs_pubinfo.find('copyrightYear')-17].strip()
			paper_type = "book"
			paper_data = {'authors':paper_authors, 'title':paper_title, 'publisher':paper_publisher, 'year':paper_date}
		elif paper_is_article and article_in_journal:
			temp_journal_data = bs_pubinfo[bs_pubinfo.find('asearch.pl'):bs_pubinfo.find('copyrightYear')-17]
			paper_journal_name = temp_journal_data[temp_journal_data.find('>')+1:temp_journal_data.rfind('a>')-2].strip()
			paper_journal_issue = temp_journal_data[temp_journal_data.find('em>')+3:temp_journal_data.rfind(':')].strip()
			paper_journal_pages = temp_journal_data[temp_journal_data.rfind(':')+1:].strip()
			paper_journal_pages = paper_journal_pages.replace(' ','')
			paper_journal_pages = paper_journal_pages.replace('--','-')
			paper_type = "article"
			paper_data = {'authors':paper_authors, 'title':paper_title, 'journal':paper_journal_name, 'issue':paper_journal_issue, 'pages':paper_journal_pages, 'year':paper_date}
		elif paper_is_article and article_in_book:
			temp_book_data = bs_pubinfo[bs_pubinfo.find('recPubInfo')+12:bs_pubinfo.find('copyrightYear')-17]
			book_authors = temp_book_data[temp_book_data.find('In')+2:temp_book_data.find('<em')].strip(', ')
			book_title = temp_book_data[temp_book_data.find('<em')+4:temp_book_data.rfind('em>')-2].strip()
			temp_book_publisher = temp_book_data[temp_book_data.rfind('em>')+4:temp_book_data.rfind('copyrightYear')-17]
			
			if temp_book_data.count(' pp. ') > 0:
				book_publisher = temp_book_data[temp_book_data.rfind('em>')+4:temp_book_data.find(' pp. ')].strip(' .')
				book_pages = temp_book_data[temp_book_data.find(' pp. ')+5:].strip(' .')
				book_pages = book_pages.replace(' ','')
				book_pages = book_pages.replace('--','-')
				paper_type = "chapter"
				paper_data = {'authors':paper_authors, 'title':paper_title, 'editors':book_authors, 'booktitle':book_title, 'publisher':book_publisher, 'pages':book_pages, 'year':paper_date}
			else:
				book_publisher = temp_book_data[temp_book_data.rfind('em>')+4:].strip('  .')
				book_pages = '(brak wyszczególnionych stron)'
				paper_type = "chapter"
				paper_data = {'authors':paper_authors, 'title':paper_title, 'editors':book_authors, 'booktitle':book_title, 'publisher':book_publisher, 'pages':book_pages, 'year':paper_date}
		
		return paper_type, paper_data


class MakeDocx:
		
	def save_document(doc_to_save):
	
		doc_to_save.save('test.docx')
		print('docx saved')
		
	def format_document(doc_to_format):
		
		# standardowe wymiary
		standard_height = Cm(29.7)
		standard_width = Cm(21)
		standard_margin = Cm(2.5)
		
		section = doc_to_format.sections[0]
		
		section.page_height = standard_height
		section.page_width = standard_width
		section.bottom_margin = standard_margin
		section.top_margin = standard_margin
		section.left_margin = standard_margin
		section.right_margin = standard_margin
	
	def iter_paragraphs(entry_list):
		
		document_to_iter = Document()
		
		for entry in entry_list:
			print('iter_paragraphs entry:')
			print(entry)
			document_to_iter = MakeDocx.add_paragraph_apa(entry[0], entry[1], document_to_iter)
			print('przerobione')
		
		MakeDocx.format_document(document_to_iter)
		MakeDocx.save_document(document_to_iter)
		print('przerob gotowy')
		
	def add_paragraph_apa(entry_type, entry_data, doc):
		
		if entry_type == 'article':
			
			entry_par = doc.add_paragraph()
			
			entry_authors = entry_par.add_run(entry_data['authors']+' ')
			entry_year = entry_par.add_run('('+entry_data['year']+'). ')
			entry_title = entry_par.add_run(unescape(entry_data['title'])+'. ')
			entry_journal = entry_par.add_run(unescape(entry_data['journal'])).italic = True
			entry_issue = entry_par.add_run(', '+entry_data['issue']+', ')
			entry_pages = entry_par.add_run(entry_data['pages']+'.')
			
			#  Nazwisko, X., Nazwisko2, X. Y., Nazwisko3, Z. (rok). Tytuł artykułu. Tytuł Czasopisma, nr rocznika(nr zeszytu), strona początku–strona końca.
		
		elif entry_type == 'book':
			
			entry_par = doc.add_paragraph()
			
			entry_authors = entry_par.add_run(entry_data['authors']+' ')
			entry_year = entry_par.add_run('('+entry_data['year']+'). ')
			entry_title = entry_par.add_run(unescape(entry_data['title'])+'. ').italic = True
			entry_publisher = entry_par.add_run(unescape(entry_data['publisher'])+'.')
			
			#  Nazwisko, X., Nazwisko, X. Y. (rok). Tytuł książki. Miejsce wydania: Wydawnictwo.
			
		elif entry_type == 'chapter':
		
			entry_par = doc.add_paragraph()
	
			entry_authors = entry_par.add_run(entry_data['authors']+' ')
			entry_year = entry_par.add_run('('+entry_data['year']+'). ')
			entry_title = entry_par.add_run(unescape(entry_data['title'])+'. W: ')
			entry_editors = entry_par.add_run(unescape(entry_data['editors'])+', ')
			entry_booktitle = entry_par.add_run(unescape(entry_data['booktitle'])).italic = True
			entry_pages = entry_par.add_run(' (s. '+entry_data['pages']+'). ')
			entry_publisher = entry_par.add_run(unescape(entry_data['publisher'])+'.')
			
			#  Nazwisko, X. (rok). Tytuł rozdziału. W: Y. Nazwisko, B. Nazwisko (red.), Tytuł książki (s. strona początku–strona końca). Miejsce wydania: Wydawnictwo.
	
		return doc
	
	
input_urls = input('urls: ')

urls_checked = Philpapers.check_urls(input_urls)

if len(urls_checked) > 0:

	philpapers_entries_list = []
	
	for item in urls_checked:
		entry_data = Philpapers.get_data(item)
		integrity_check = Philpapers.check_integrity(entry_data, item)
		print(integrity_check)
		if integrity_check:
			philpapers_entries_list.append(Philpapers.parse_data(entry_data))
		else:
			exception_reason = 'Nie można pobrać danych z %s : odnośnik nie zawiera poprawnego schematu strony Philpapers (brak typu dokumentu). Popraw listę adresów i spróbuj jeszcze raz' % (item)
			Exceptions.signal_exception(exception_reason)

	# MakeDocx.iter_paragraphs(philpapers_entries_list)
	return dict(display=entry_data)
	
	
# przykładowa lista url:
# https://philpapers.org/rec/BAUPAT-3 https://philpapers.org/rec/SURMCA https://philpapers.org/rec/RORTGO https://philpapers.org/rec/OPPACT https://philpapers.org/rec/SMATTT https://philpapers.org/rec/RORTGO